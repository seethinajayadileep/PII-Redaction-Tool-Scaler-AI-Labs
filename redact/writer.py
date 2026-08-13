"""Write redacted pages into a .docx file."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def write_docx(pages: list[str], destination: str | Path, *, source_name: str, counts: dict[str, int]) -> Path:
    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = doc.add_heading('Redacted document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph()
    run = intro.add_run(f'Source: {source_name}. Personally identifiable information was replaced with consistent fake stand-ins. This file is not the original prospectus.')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(68, 68, 68)
    if counts:
        summary = doc.add_paragraph('Replacements in this run: ')
        summary.add_run(', '.join((f'{pii_type}={n}' for pii_type, n in sorted(counts.items()))))
    for index, page in enumerate(pages, start=1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f'Page {index}', level=2)
        chunks = [part.strip() for part in (page or '').split('\n') if part.strip()]
        if not chunks:
            body = doc.add_paragraph('[empty page]')
            body.runs[0].font.size = Pt(10)
            continue
        for chunk in chunks:
            body = doc.add_paragraph(chunk)
            if body.runs:
                body.runs[0].font.size = Pt(10)
    doc.save(destination)
    return destination